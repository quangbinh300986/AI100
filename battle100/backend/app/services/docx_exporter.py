import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, font_name="微软雅黑", size_pt=10.5, bold=False, color_rgb=None):
    """设置 Word 段落块的字体样式"""
    run.font.name = font_name
    # 针对 Word 软件内置的中文字体映射进行 oxml 特殊属性处理
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb


def set_cell_background(cell, hex_color):
    """设置表格单元格的背景颜色 (16进制值)"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_formatted_text(p, text_str):
    """解析 Markdown 语法的 **加粗** 文字并渲染段落"""
    parts = re.split(r'(\*\*.*?\*\*)', text_str)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = p.add_run(bold_text)
            # 加粗文本高亮显示为深蓝蓝色，醒目排版
            set_font(run, font_name="微软雅黑", size_pt=10.5, bold=True, color_rgb=RGBColor(16, 42, 76))
        else:
            run = p.add_run(part)
            set_font(run, font_name="微软雅黑", size_pt=10.5, bold=False)


def add_metrics_table(doc, metrics):
    """插入浅灰底 Word 3x3 表格，展示团队核心业绩指标"""
    items = [
        ("🏆 营销新签合同额", f"{metrics.get('marketing_signed', 0.0):.2f} 万元"),
        ("📋 交付新签合同额", f"{metrics.get('delivery_signed', 0.0):.2f} 万元"),
        ("🎯 中标项目个数", f"{metrics.get('win_bids', 0)} 个"),
        ("😊 幸福动作个数", f"{metrics.get('happiness_count', 0)} 次"),
        ("🤝 铁三角联动次数", f"{metrics.get('triangle_count', 0)} 次"),
        ("📌 有效商机线索量", f"{metrics.get('valid_leads', 0)} 个"),
        ("🔥 潜力商机线索量", f"{metrics.get('potential_leads', 0)} 个"),
        ("💰 CRM 累计确认产值", f"{metrics.get('production_value', 0.0):.2f} 万元"),
        ("💵 CRM 到账回款额", f"{metrics.get('receive_value', 0.0):.2f} 万元")
    ]
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # 填充表格数据并调整外观格式
    for idx, (label, val) in enumerate(items):
        row_idx = idx // 3
        col_idx = idx % 3
        cell = table.cell(row_idx, col_idx)
        
        # 垂直居中设置，虽然默认靠上，但增加段落前后的 padding 也可以很规整
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        # 指标名称
        run_label = p.add_run(f"{label}\n")
        set_font(run_label, font_name="微软雅黑", size_pt=9.5, bold=True, color_rgb=RGBColor(89, 89, 89))
        
        # 指标数值
        run_val = p.add_run(val)
        set_font(run_val, font_name="微软雅黑", size_pt=11.5, bold=True, color_rgb=RGBColor(16, 42, 76))
        
        # 设置单元格的浅灰色背景
        set_cell_background(cell, "F2F2F2")


def parse_markdown_to_docx(doc, content):
    """解析 Markdown 各级大纲、列表、粗体并转换为 Word 排版段落"""
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # 1. 过滤 Markdown 横线
        if stripped == '---' or stripped == '---  ':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("—" * 40)
            set_font(run, font_name="微软雅黑", size_pt=10, bold=False, color_rgb=RGBColor(217, 217, 217))
            continue
            
        # 2. 一级标题
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            set_font(run, font_name="微软雅黑", size_pt=16, bold=True, color_rgb=RGBColor(16, 42, 76))
            continue
            
        # 3. 二级标题
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[3:].strip())
            set_font(run, font_name="微软雅黑", size_pt=13.5, bold=True, color_rgb=RGBColor(24, 144, 255))
            continue
            
        # 4. 三级标题
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:].strip())
            set_font(run, font_name="微软雅黑", size_pt=11.5, bold=True, color_rgb=RGBColor(89, 89, 89))
            continue
            
        # 5. 无序列表 (识别 - 或 * 且支持二级缩进)
        is_list = False
        level = 0
        list_prefix = ""
        
        # 匹配嵌套的缩进列表 (如 "  - 子项")
        if line.startswith("  - ") or line.startswith("    - ") or line.startswith("  * ") or line.startswith("    * "):
            is_list = True
            level = 1
            list_prefix = "  •  "
            content_text = stripped[2:].strip()
        elif stripped.startswith("- ") or stripped.startswith("* "):
            is_list = True
            level = 0
            list_prefix = "•  "
            content_text = stripped[2:].strip()
            
        if is_list:
            p = doc.add_paragraph()
            # 设置悬挂缩进和段后空间
            p.paragraph_format.left_indent = Pt(20 * (level + 1))
            p.paragraph_format.space_after = Pt(3)
            
            run_prefix = p.add_run(list_prefix)
            set_font(run_prefix, font_name="微软雅黑", size_pt=10.5, bold=True)
            
            add_formatted_text(p, content_text)
            continue
            
        # 6. 普通正文段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        # 若是引用块
        if stripped.startswith(">"):
            run_quote = p.add_run("“ ")
            set_font(run_quote, font_name="微软雅黑", size_pt=10.5, bold=True, color_rgb=RGBColor(128, 128, 128))
            
            body_text = stripped[1:].strip()
            if body_text.startswith("*") and body_text.endswith("*"):
                body_text = body_text[1:-1].strip() # 剥离斜体字
                
            add_formatted_text(p, body_text)
            
            run_quote_end = p.add_run(" ”")
            set_font(run_quote_end, font_name="微软雅黑", size_pt=10.5, bold=True, color_rgb=RGBColor(128, 128, 128))
        else:
            add_formatted_text(p, stripped)


def export_markdown_to_docx(title: str, metrics: dict | None, content: str) -> BytesIO:
    """提供通用的 Markdown 周报转化为 Word (.docx) 的方法，返回 BytesIO 二进制流"""
    doc = Document()
    
    # 调整 Word 页面边距，使排版紧凑
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)
        
    # 1. 插入大标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(title)
    set_font(run_title, font_name="微软雅黑", size_pt=20, bold=True, color_rgb=RGBColor(16, 42, 76))
    
    # 2. 插入指标表格
    if metrics:
        add_metrics_table(doc, metrics)
        # 增加表格和下方正文的间距段落
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(6)
        p_space.paragraph_format.space_after = Pt(6)
        
    # 3. 转换 Markdown 正文
    parse_markdown_to_docx(doc, content)
    
    # 保存到内存中并返回
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
